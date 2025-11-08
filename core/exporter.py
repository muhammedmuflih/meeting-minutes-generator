# core/exporter.py
from docx import Document
from docx.shared import Inches as DocxInches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
import os
from datetime import datetime

def export_to_text(minutes_data, output_filepath):
    """Exports minutes to a plain text file."""
    try:
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write("--- Meeting Minutes ---\n\n")
            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

            f.write("Meeting Summary:\n")
            f.write(minutes_data.get('summary', 'N/A') + "\n\n")

            f.write("Key Decisions:\n")
            f.write(minutes_data.get('decisions', 'N/A') + "\n\n")

            f.write("Action Items:\n")
            f.write(minutes_data.get('action_items', 'N/A') + "\n\n")

            f.write("Important Deadlines:\n")
            f.write(minutes_data.get('deadlines', 'N/A') + "\n\n")

        print(f"Minutes exported to {output_filepath}")
        return output_filepath
    except Exception as e:
        print(f"Error exporting to text: {e}")
        raise

def export_to_word(minutes_data, output_filepath):
    """Exports minutes to a Word (.docx) file."""
    try:
        document = Document()

        # Title
        document.add_heading('Meeting Minutes', level=0)
        document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        document.add_paragraph('') # Empty paragraph for spacing

        # Summary
        document.add_heading('1. Meeting Summary', level=1)
        document.add_paragraph(minutes_data.get('summary', 'N/A'))
        document.add_paragraph('')

        # Decisions
        document.add_heading('2. Key Decisions', level=1)
        document.add_paragraph(minutes_data.get('decisions', 'N/A'))
        document.add_paragraph('')

        # Action Items
        document.add_heading('3. Action Items', level=1)
        document.add_paragraph(minutes_data.get('action_items', 'N/A'))
        document.add_paragraph('')

        # Deadlines
        document.add_heading('4. Important Deadlines', level=1)
        document.add_paragraph(minutes_data.get('deadlines', 'N/A'))
        document.add_paragraph('')

        document.save(output_filepath)
        print(f"Minutes exported to {output_filepath}")
        return output_filepath
    except Exception as e:
        print(f"Error exporting to Word: {e}")
        raise

def export_to_pdf(minutes_data, output_filepath):
    """Exports minutes to a PDF file using ReportLab."""
    try:
        doc = SimpleDocTemplate(output_filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        h0_style = ParagraphStyle(name='Title', fontSize=24, leading=28, alignment=TA_CENTER,
                                  fontName='Helvetica-Bold')
        h1_style = ParagraphStyle(name='Heading1', fontSize=18, leading=22, alignment=TA_LEFT,
                                  fontName='Helvetica-Bold')
        normal_style = ParagraphStyle(name='Normal', fontSize=10, leading=12, alignment=TA_LEFT,
                                      fontName='Helvetica')

        # Title
        story.append(Paragraph("Meeting Minutes", h0_style))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))

        # Summary
        story.append(Paragraph("1. Meeting Summary", h1_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(minutes_data.get('summary', 'N/A'), normal_style))
        story.append(Spacer(1, 0.2 * inch))

        # Decisions
        story.append(Paragraph("2. Key Decisions", h1_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(minutes_data.get('decisions', 'N/A'), normal_style))
        story.append(Spacer(1, 0.2 * inch))

        # Action Items
        story.append(Paragraph("3. Action Items", h1_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(minutes_data.get('action_items', 'N/A'), normal_style))
        story.append(Spacer(1, 0.2 * inch))

        # Deadlines
        story.append(Paragraph("4. Important Deadlines", h1_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(minutes_data.get('deadlines', 'N/A'), normal_style))
        story.append(Spacer(1, 0.3 * inch))

        doc.build(story)
        print(f"Minutes exported to {output_filepath}")
        return output_filepath
    except Exception as e:
        print(f"Error exporting to PDF: {e}")
        raise

if __name__ == '__main__':
    # Example usage for testing
    sample_minutes = {
        "summary": "The team discussed the Q3 marketing strategy, focusing on social media platforms. Key decisions were made regarding platform priorities and budget allocation.",
        "decisions": "We decided to proceed with the green project. New marketing budget approved.",
        "action_items": "John to create presentation slides by Tuesday. Sarah to follow up with vendor. Mark to update budget spreadsheet.",
        "deadlines": "Presentation slides due: Tuesday. Vendor follow-up: End of week. Budget update: Next Monday."
    }

    output_dir = "../outputs"
    os.makedirs(output_dir, exist_ok=True)
    file_prefix = datetime.now().strftime("%Y%m%d_%H%M%S_meeting_minutes")

    # Export to Text
    text_path = os.path.join(output_dir, f"{file_prefix}.txt")
    export_to_text(sample_minutes, text_path)

    # Export to Word
    word_path = os.path.join(output_dir, f"{file_prefix}.docx")
    export_to_word(sample_minutes, word_path)

    # Export to PDF
    pdf_path = os.path.join(output_dir, f"{file_prefix}.pdf")
    export_to_pdf(sample_minutes, pdf_path)